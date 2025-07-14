import os 
import re
import json
import random
import requests
import datetime
import shutil
import numpy as np #1.26.4
import torch #2.2.0+cu121
import pandas as pd #2.2.0
import torch.nn as nn 
import torch.nn.functional as F
import torch.optim as optim
import openpyxl
import csv
from ltp import LTP #4.2.13
from graphviz import Digraph #0.20.1
from docx import Document #1.1.0
from transformers import BertTokenizer, BertForSequenceClassification, BertConfig #4.38.2
from sklearn.model_selection import train_test_split
from transformers import AdamW
from torch.utils.data import Dataset
from torch.utils.data import DataLoader
global_model_version = 0
original_string = """农、林、牧、渔业
采矿业
制造业
电力、热力、燃气及水生产和供应业
建筑业
批发和零售业
交通运输、仓储和邮政业
住宿和餐饮业
信息传输、软件和信息技术服务业
金融业
房地产业
租赁和商务服务业
科学研究和技术服务业
水利、环境和公共设施管理业
居民服务、修理和其他服务业
教育
卫生和社会工作
文化、体育和娱乐业
无法确定"""
business_natures = original_string.split('\n')
def downloadreport_year(year,path = '.',number=10):
    path=path+"/"
    year=int(year)
    year+=1
    yearr=str(year-1)

    folder_name = os.path.join(path, yearr)
    if not os.path.exists(folder_name):
        os.makedirs(folder_name)

    start_date = datetime.date(year=year, month=1, day=1)
    end_date = datetime.date(year=year, month=4, day=30)
    year=str(year)
    num=0
    listt=[]
    current_date = start_date
    while current_date <= end_date:
        month = current_date.month
        day = current_date.day
        response = requests.get(
            'http://query.sse.com.cn/infodisplay/queryLatestBulletinNew.do?jsonCallBack=jsonpCallback49000&isPagination=true&productId=&keyWord=&reportType2=DQGG&reportType=YEARLY&beginDate={}-01-01&endDate={}-{}-{}&pageHelp.pageSize=25&pageHelp.pageCount=50&pageHelp.pageNo=1&pageHelp.beginPage=1&pageHelp.cacheSize=1&pageHelp.endPage=5&_=1704038400000'.format(year,year,month,day),
            headers={'Referer': 'http://www.sse.com.cn/disclosure/listedinfo/announcement/'}
        )

        json_str = response.text[19:-1]
        data = json.loads(json_str)

        for report in data['result']:
            if "摘要" in report['title']:
                continue
            if report['title'] in listt:
                continue
            num=num+1
            listt.append(report['title'])

            download_url = 'http://static.sse.com.cn' + report['URL']
            ccode = download_url
            file_name = report['title'] +ccode[77:83]+ '.pdf'
            print(download_url,file_name)
            resource = requests.get(download_url, stream=True)

            if num<=number :
                file_path = os.path.join(path,yearr,file_name)
                with open(file_path, 'wb') as fd:
                    for chunk in resource.iter_content(102400):
                            fd.write(chunk)
                    print(file_name, '完成下载',num)
            else:
                return 0
        current_date += datetime.timedelta(days=1)


def Conv_docx2tables_by2parts(filepath):
    df_list_for_judge = []
    df_list_wait_combine = []
    for filename in os.listdir(filepath):
        doc = Document(f'{filepath}/{filename}')  # 替换为你的文件名
        dataframes_for_judge = []  # 存储所有表格的DataFrame列表
        dataframes_wait_combine = []
        elements = []  # 存储文档的所有元素
        # 遍历文档中的所有元素，收集段落和表格
        table_counter = 0
        for element in doc.element.body:
            if element.tag.endswith('p'):
                elements.append(('paragraph', element.text.strip()))
            elif element.tag.endswith('tbl'):
                elements.append(('table', doc.tables[table_counter]))
                table_counter+=1
        # 处理每个表格
        for i, (etype, content) in enumerate(elements):
            if etype == 'table':
                # 初始化用于存储当前表格数据的列表
                table_data_tbval = []
                table_data_paraval = []
                # 初始化用于存储将添加到表格顶部的段落文本的列表
                preceding_paragraphs = []
                # 向上回溯查找直到找到三个非空文本段落或遇到一个表格
                count = 0  # 用于计数找到的非空文本段落
                for prev_index in range(i - 1, -1, -1):  # 从当前位置向上遍历
                    prev_element = elements[prev_index]
                    if prev_element[0] == 'paragraph' and prev_element[1]:  # 是非空段落
                        preceding_paragraphs.insert(0, prev_element[1])  # 将段落文本插入列表的开始

                        count += 1
                        if count == 5:  # 如果已找到三个非空文本段落，则停止
                            break
                    elif prev_element[0] == 'table':  # 如果遇到另一个表格，则停止
                        if count == 0:
                            preceding_paragraphs.insert(0, '无上信息')
                        break

                # 将找到的段落作为新行添加到表格顶部
                if preceding_paragraphs[0] == '无上信息':
                    paragraph_row = [preceding_paragraphs[0]] + [None]*(len(content.rows[0].cells)-1)  # 填充段落文本到整行
                    table_data_tbval.append(paragraph_row)

                # 添加表格的其他行
                for row in content.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data_tbval.append(row_data)

                # 将找到的段落作为新行添加到表格顶部
                for para_text in preceding_paragraphs:
                    paragraph_row = [para_text] + [None]*(len(content.rows[0].cells)-1)  # 填充段落文本到整行
                    table_data_paraval.append(paragraph_row)

                # 添加表格的第一行
                for row in content.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data_paraval.append(row_data)
                    break

                # 将处理好的表格数据转换成DataFrame并存储
                df_paraval = pd.DataFrame(table_data_paraval)
                df_tbval = pd.DataFrame(table_data_tbval)

                dataframes_for_judge.append(df_paraval)
                dataframes_wait_combine.append(df_tbval)
                
        df_list_for_judge.append(dataframes_for_judge)   
        df_list_wait_combine.append(dataframes_wait_combine)

    #把无上信息的那种表格去掉, 这些不判断, 浪费资源
    df_list_for_judge_without_up_info = []
    index_of_df_list_for_judge_without_up_info = []
    for df_with_no_up_info in df_list_for_judge:
        df_company_without_up_info = []
        index_of_company_df_list_for_judge_without_up_info = []
        for i, df in enumerate(df_with_no_up_info):
            if df.iloc[0,0] != '无上信息':
                df_company_without_up_info.append(df)
                index_of_company_df_list_for_judge_without_up_info.append(i)
        df_list_for_judge_without_up_info.append(df_company_without_up_info)
        index_of_df_list_for_judge_without_up_info.append(index_of_company_df_list_for_judge_without_up_info)
    return df_list_for_judge_without_up_info, index_of_df_list_for_judge_without_up_info, df_list_wait_combine

#清洗文本函数
def remove_punctuation(text):
    # 使用正则表达式替换掉文本中的标点符号
    punctuation = """ )(！？｡。＂＃＄％＆＇（）＊＋，－／：；＜＝＞＠［＼］＾＿｀｛｜｝～｟｠｢｣､、〃《》「」『』【】〔〕〖〗〘〙〚〛〜〝□〞〟〰〾〿√–—‘’‛“”„‟…‧﹏.\n\r\t"""
    return re.sub(f"[{punctuation}]", "", text)

# 将表格反转转换为文本
def table_to_text_for_one_report(table):
    train_text_without_none = [x for x in [' '.join(str(cell) for cell in row[::-1] if pd.isnull(cell) != True) for row in table.values[::-1]]]
    return ' '.join(train_text_without_none)

def convert_tables_to_tensor(df_list_for_judge, labels_without_up_info):
    text_sequences_list = []
    indices_af_del = []
    for i ,report_tb_labels in enumerate(labels_without_up_info):
        report_tables = df_list_for_judge[i]
        # 准备数据
        all_texts = [table_to_text_for_one_report(data) for data in report_tables]
        # 要查找的字符
        chars_to_check = ["实际控制人", "股东",'子公司']
        # 使用列表推导式检查每个元素是否包含任何一个指定字符，并将包含字符的元素添加到新列表中
        all_texts_af_del = [item for item in all_texts if any(char in item for char in chars_to_check)]
        # 使用列表推导式获取包含指定字符的元素的索引
        indices = [report_tb_labels[index] for index, item in enumerate(all_texts) if any(char in item for char in chars_to_check)]
        clean_texts = [remove_punctuation(text) for text in all_texts_af_del]
        text_sequences_list.append(clean_texts)
        indices_af_del.append(indices)
    return text_sequences_list, indices_af_del


def predict_fun(df_list_for_judge_without_up_info, index_of_df_list_for_judge_without_up_info):
    device = torch.device("cuda" if torch.cuda.is_available() else "cpu")

    # 载入模型
    # 定义 tokenizer，传入词汇表
    tokenizer = BertTokenizer('models_find_tables/vocab.txt')

    # 超参数
    hidden_dropout_prob = 0.3
    num_labels = 4
    config = BertConfig.from_pretrained("bert-base-uncased", num_labels=num_labels, hidden_dropout_prob=hidden_dropout_prob)
    
    model = BertForSequenceClassification(config=config)
    model.load_state_dict(torch.load('models_find_tables/model_find_tables.pth', map_location=torch.device('cpu')))
    model.to(device)  # 将模型移动到指定设备上
    model.eval()  # 设置为评估模式
    
    max_len = 100  # 或其他适合你模型的长度
    test_sequences_list, test_indices_af_del = convert_tables_to_tensor(df_list_for_judge_without_up_info, index_of_df_list_for_judge_without_up_info)
    # 存储预测结果
    predictions_of_all_comp = []
    # 批量处理和预测
    for diffcompanydata in test_sequences_list:
        # 预处理文本
        inputs = tokenizer(diffcompanydata, max_length=max_len, padding=True, truncation=True, return_tensors="pt")
        inputs = {k: v.to(device) for k, v in inputs.items()}  # 将数据移至GPU

        # 进行预测
        with torch.no_grad():  # 不追踪梯度
            outputs = model(**inputs)
            predictions = outputs.logits

        # 将多出来的类表格, 进行再比较, 得到分最高的, 作为那个唯一表格, 此处是唯一的股东表, 子公司表, 实控人表
        # 原始分类，假设是通过模型输出得分的最大值决定
        original_classes = torch.argmax(predictions, dim=1)

        # 新分类，初始设为原始分类
        new_classes = original_classes.clone()

        # 对于每一个应该是唯一的类别
        for target_class in range(3):  # 对于前三个类别
            # 找出被错误分类到该类别的所有表格索引
            classified_indices = (original_classes == target_class).nonzero(as_tuple=True)[0]

            # 如果超过一个表格被分类为该类别
            if len(classified_indices) > 1:
                # 获取这些表格在该类别上的得分
                scores = predictions[classified_indices, target_class]
                # 找到得分最高的表格
                highest_score_index = classified_indices[torch.argmax(scores)]
                # 重新分类得分较低的表格为第四类
                for index in classified_indices:
                    if index != highest_score_index:
                        new_classes[index] = 3  # 重新分类为“无用的表格”
        predictions_of_all_comp.append(new_classes.cpu().numpy().tolist())
    return predictions_of_all_comp, test_indices_af_del

def predict_fun_for_business_nature(texts):
    business_natures = original_string.split('\n')
    device = torch.device("cuda" if torch.cuda.is_available() else "cpu")

    # 载入模型
    # 定义 tokenizer，传入词汇表
    tokenizer = BertTokenizer('model_business_classification/vocab.txt')

    # 超参数
    hidden_dropout_prob = 0.3
    num_labels = 19
    config = BertConfig.from_pretrained("bert-base-uncased", num_labels=num_labels, hidden_dropout_prob=hidden_dropout_prob)

    model = BertForSequenceClassification(config=config)
    model.load_state_dict(torch.load('model_business_classification/model_business_classification4.pth', map_location=torch.device('cpu')))
    model.to(device)  # 将模型移动到指定设备上
    model.eval()  # 设置为评估模式

    max_len =  10 # 或其他适合你模型的长度

    # 预处理文本
    inputs = tokenizer(texts.values.tolist(), max_length=max_len, padding=True, truncation=True, return_tensors="pt")
    inputs = {k: v.to(device) for k, v in inputs.items()}  

    # 进行预测
    with torch.no_grad():  # 不追踪梯度
        outputs = model(**inputs)
        predictions = outputs.logits

    original_classes = torch.argmax(predictions, dim=1)
    selected_business_nature = [business_natures[i] for i in original_classes]
    return selected_business_nature

def saveas_xlsx_by_prediction(test_file_path, df_list_wait_combine, predictions_of_all_comp, test_indices_af_del):
    #判断文件夹是否存在,不存在则重新创建
    if not os.path.exists('xlsx被模型标记完成的'):
        os.makedirs('xlsx被模型标记完成的')
    for i, filename in enumerate(os.listdir(test_file_path)):
        table_df_of_diffcomp = df_list_wait_combine[i]
        predictions_of_one_comp = predictions_of_all_comp[i]
        corr_index_with_predictions_of_one_comp = test_indices_af_del[i]
        corr_dict = {}
        state_list = [0,0,0]
        for i, pre_outcome in enumerate(predictions_of_one_comp):
            if pre_outcome == 0:
                corr_dict['股东表'] = corr_index_with_predictions_of_one_comp[i]
                state_list[0] = 1
            elif pre_outcome == 1:
                corr_dict['子公司表'] = corr_index_with_predictions_of_one_comp[i]
                state_list[1] = 1
            elif pre_outcome == 2:
                corr_dict['实际控制人表'] = corr_index_with_predictions_of_one_comp[i]
                state_list[2] = 1
            else:
                continue
        # 使用ExcelWriter保存多个DataFrame到一个Excel文件中
        with pd.ExcelWriter(f'xlsx被模型标记完成的/{filename[0:-5]}.xlsx') as writer:
            for i, df in enumerate(table_df_of_diffcomp):
                # 如果是特定索引的表格，使用特定的sheet名
                if state_list[0] == 1 and i == corr_dict['股东表']:
                    df.to_excel(writer, sheet_name='股东表', index=False)
                elif state_list[1] == 1 and i == corr_dict['子公司表']:
                    df.to_excel(writer, sheet_name='子公司表', index=False)
                elif state_list[2] == 1 and i == corr_dict['实际控制人表']:
                    df.to_excel(writer, sheet_name='实际控制人表', index=False)
                else:
                    # 否则，使用默认的sheet名，如'Sheet1', 'Sheet2', ...
                    df.to_excel(writer, sheet_name=f'Sheet{i + 1}', index=False)
def Judge_fun(file_path):
    df_willbe_judged, index_of_df_be_judged, df_willbe_combined = Conv_docx2tables_by2parts(file_path)
    predictions_of_all_comp, test_indices_af_del = predict_fun(df_willbe_judged, index_of_df_be_judged)
    saveas_xlsx_by_prediction(file_path, df_willbe_combined, predictions_of_all_comp, test_indices_af_del)






#draw_figure_code
def concat_tables(filepath):

    filderpath = 'xlsx包括合成好的三张表'
    os.makedirs(filderpath, exist_ok=True)  # 创建文件夹
    for filename in os.listdir(filepath):
        try:
            # 用 pd.ExcelFile() 创建 ExcelFile 对象
            excel_file = pd.ExcelFile(f'xlsx被模型标记完成的/{filename}')
            # 获取所有 sheet 的名称
            sheet_names = excel_file.sheet_names
            share_holder_table = pd.DataFrame()
            composedcompanys_table = pd.DataFrame()
            controller_table = pd.DataFrame()

            if '股东表' in sheet_names:
                share_holder = pd.read_excel(excel_file, sheet_name='股东表',header=0)
                for i in range(sheet_names.index('股东表'), len(sheet_names)):
                    Next_table = pd.read_excel(excel_file, sheet_name=sheet_names[i+1],header=0)
                    if Next_table.iloc[0, 0] != '无上信息':
                        break
                    share_holder = pd.concat([share_holder, Next_table.iloc[1:,:]], ignore_index=True)
                share_holder_table = share_holder

            if '子公司表' in sheet_names:
                composedcompanys = pd.read_excel(excel_file, sheet_name='子公司表',header=0)
                for i in range(sheet_names.index('子公司表'), len(sheet_names)):
                    Next_table = pd.read_excel(excel_file, sheet_name=sheet_names[i+1],header=0)
                    if Next_table.iloc[0, 0] != '无上信息':
                        break
                    composedcompanys = pd.concat([composedcompanys, Next_table.iloc[1:, :]], ignore_index=True)
                composedcompanys_table = composedcompanys

            if '实际控制人表' in sheet_names:
                controller = pd.read_excel(excel_file, sheet_name='实际控制人表',header=0)
                for i in range(sheet_names.index('实际控制人表'), len(sheet_names)):
                    Next_table = pd.read_excel(excel_file, sheet_name=sheet_names[i+1],header=0)
                    if Next_table.iloc[0, 0] != '无上信息':
                        break
                    controller = pd.concat([controller, Next_table.iloc[1:, :]], ignore_index=True)
                controller_table = controller

            # 将三个 DataFrame 分别写入不同的工作表

            with pd.ExcelWriter(f'{filderpath}/{filename[0:-5]}.xlsx') as writer:
                share_holder_table.to_excel(writer,sheet_name='股东', index=False)
                composedcompanys_table.to_excel(writer,sheet_name='子公司', index=False)
                controller_table.to_excel(writer, sheet_name='实控人', index=False)
        except Exception as e:
            print(f"(failed to concat tables of {filename})Error: {e}")
# def del_useless_head_lines(df, keyword):
#     #处理表的多余表头
#     index_after_drop = 0.1
#     df = df.replace([' ','\t'], '', regex=True).replace('\n', '', regex=True)
#     for i in range(len(df)):
#         if keyword in df.iloc[i,0]:
#             if keyword in df.iloc[i+1,0]:
#                 index_after_drop = i
#                 break
#             else:
#                 continue
#         else:
#             continue
#     if index_after_drop != 0.1:
#         df = df.loc[index_after_drop+1:,:]
#     df = df.reset_index(drop=True)
#     new_header = df.iloc[0]
#     new_header = new_header.fillna('未知')
#     df = df[1:]
#     df.columns = new_header
#     df.reset_index(drop=True,inplace=True)
#     return df
            
def del_useless_head_lines(df, keyword):
    #处理表的多余表头
    index_after_drop = 0.1
    df = df.replace([' ','\t'], '', regex=True).replace('\n', '', regex=True)
    for i in range(len(df)):
        if keyword in str(df.iloc[i,0]):
            if keyword not in str(df.iloc[i+1,0]):
                index_after_drop = i
                break
            else:
                continue
        else:
            continue
    if index_after_drop != 0.1:
        df = df.loc[index_after_drop:,:]

    # 获取第一列的列名
    first_column = df.columns[0]
    cleaned_df = df.dropna(subset=[first_column])
    cleaned_df = cleaned_df.reset_index(drop=True)
    new_header = cleaned_df.iloc[0]
    new_header = new_header.fillna('未知')
    cleaned_df = cleaned_df[1:]
    cleaned_df.columns = new_header
    cleaned_df.reset_index(drop=True,inplace=True)
    return cleaned_df

def merge_rows(df):
    for i in range(len(df)):
        if pd.isnull(df.iloc[i, :]).any():
            for j in range(len(df.columns)):
                df.iloc[i-1, j] = f"{df.iloc[i-1, j]} {df.iloc[i, j]}".strip()
    # 删除那些已经被合并的行
    return df.dropna().reset_index(drop=True)

def process_composedcompanys_table(df):
    #处理公司性质的空值
    column_name_nature_chara = [col for col in df.columns if '业务' in col]
    for i in range(2, len(df)-1):
        df[column_name_nature_chara[0]] = df[column_name_nature_chara[0]].fillna('其他')
    punctuation = '-;:,.!'  
    # 使用 list 函数将字符串转换为列表，并在 df.replace 中使用
    df = df.fillna('0')
    df = merge_rows(df)
    df = df.replace(' ', '', regex=True)
    column_name_direct = [col for col in df.columns if '直接' in col]
    column_name_indirect = [col for col in df.columns if '间接' in col]
    if column_name_direct or column_name_indirect:
        # 去除百分比符号，并转换为浮点数
        df[column_name_direct[0]] = df[column_name_direct[0]].replace('%', '', regex=True)
        df[column_name_direct[0]] = df[column_name_direct[0]].replace('-', None).astype(float)
        df[column_name_indirect[0]] = df[column_name_indirect[0]].replace('%', '', regex=True)
        df[column_name_indirect[0]] = df[column_name_indirect[0]].replace('-', np.nan).astype(float)
        df = df.fillna(0)
        # 合并两列
        df[column_name_direct[0]] = df[column_name_direct[0]] + df[column_name_indirect[0]]
        df.drop(columns=[column_name_indirect[0]], inplace=True)
    else:
        if len(column_name_direct) != 0:
            df[column_name_direct[0]] = df[column_name_direct[0]].str.replace('%', '').astype(float)
    # 指定你需要保留的列，并按顺序排列
    selected_columns_dict = {}
    for column_name in df.columns:
        if '子公司名称' in column_name:
            selected_columns_dict['子公司名称'] = column_name
        elif '主要经营地' in column_name:
            selected_columns_dict['主要经营地'] = column_name
        elif '注册地' in column_name:
            selected_columns_dict['注册地'] = column_name
        elif '业务' in column_name:
            selected_columns_dict['业务'] = column_name
        elif '直接' in column_name:
            selected_columns_dict['直接'] = column_name
        elif '方式' in column_name:
            selected_columns_dict['取得'] = column_name
    # 使用 loc 方法选择这些列，并赋值给一个新的 DataFrame
    selected_columns = [selected_columns_dict['子公司名称'],
                        selected_columns_dict['主要经营地'],
                        selected_columns_dict['注册地'],
                        selected_columns_dict['业务'],
                        selected_columns_dict['直接'],
                        selected_columns_dict['取得']]
    new_df = df.loc[:, selected_columns]
    return new_df

def process_shareholder_table(df):
    column_shareholdername = [col for col in df.columns if '名称' in col]
    # share_holder_nature_column = [col for col in df.columns if '性质' in col]
    contains_word = df.apply(lambda row: row.astype(str).str.contains('境外').any(), axis=1)
    df['股东性质'] = np.where(contains_word, '境外', '境内')
    
    column_ratio = [col for col in df.columns if '比例' in col]
    df = df[[column_shareholdername[0],'股东性质',column_ratio[0]]]
    df = df.iloc[0:10]
    return df



# filename = '603868_20240312_29T7.xlsx'
# filepath = f'xlsx包括合成好的三张表/{filename}'
#实际控制人表格读取
def get_three_tables_and_namelist(filename, filepath):
    share_holder_df = pd.read_excel(filepath, sheet_name='股东',header=0)
    composed_companys_df = pd.read_excel(filepath, sheet_name='子公司',header=0)
    controller_df = pd.read_excel(filepath, sheet_name='实控人',header=0)

    # 提取关键词之间的文本
    result_text = ''.join(controller_df.values.flatten().astype(str))

    #处理子公司表
    composed_companys_df = del_useless_head_lines(composed_companys_df,'子公司名称')
    composed_companys_df = process_composedcompanys_table(composed_companys_df)
    predicted_business_natures = predict_fun_for_business_nature(composed_companys_df.iloc[:,3])
    composed_companys_df.iloc[:,3] = predicted_business_natures
    #处理股东表
    share_holder_df = del_useless_head_lines(share_holder_df, '股东名称')
    share_holder_df = process_shareholder_table(share_holder_df)

    filepath = f'xlsx包括合成好的三张表/{filename}'
    #实际控制人表格读取
    controller_df = pd.read_excel(filepath, sheet_name='实控人')
    # Check if the DataFrame is empty
    name_list = []
    if controller_df.empty == False:
        # 提取关键词之间的文本
        result_text = ''.join(controller_df.values.flatten().astype(str))
        # 创建LTP的pipeline
        ltp = LTP()
        # 对句子进行分词和命名实体识别
        seg_result = ltp(result_text)
        # 命名实体识别结果
        ner_result = seg_result[2]
        #定义一个人名集合
        name_set = set()
        # 输出人名
        for ner in ner_result:
            tag, name = ner
            if tag == 'Nh':
                name = name.strip()
                name_set.add(name)
        name_list = list(name_set)
        if len(name_list) == 0:
            name_list.append('未能识别实际控制人')
    else:
        name_list.append('未能识别实际控制人')
    data_list_after_combine = []
    data_list_after_combine.append(share_holder_df)
    data_list_after_combine.append(composed_companys_df)

    #得到不重复的子公司性质series
    Subsidiary_characteristics = data_list_after_combine[1].iloc[:,3]
    Subsidiary_characteristics.drop_duplicates(keep='first',inplace=True)
    Subsidiary_characteristics.reset_index(drop=True, inplace=True)
    return data_list_after_combine, Subsidiary_characteristics, name_list



# 定义一个函数来每四个字符后添加换行符
def insert_newline_every_2_chars(s):
    return '\n'.join(s[i:i+2] for i in range(0, len(s), 2))

def insert_newline_every_4_chars(s):
    return '\n'.join(s[i:i+4] for i in range(0, len(s), 4))

#处理需要预测的数据, 并且转换为索引, 如果遇到不存在则用<UNK>代替
def text_to_indices(text, char_dict, unk_token="[UNK]"):
    indices = []
    for char in text:
        # 如果字符不在字典中，使用未知字符代替
        index = char_dict.get(char, char_dict[unk_token])
        indices.append(index)
    return indices

def draw_svg(data_list_after_combine, Subsidiary_characteristics, name_list, filename):
    dot = Digraph(comment='Organization Chart',engine='dot')

    dot.attr('node', shape='box')
    dot.attr(splines="ortho")
    dot.attr(rankdir='TB')
    # 设置图级别的属性
    dot.graph_attr['ranksep'] = '1'
    dot.graph_attr['nodesep'] = '4.4'
    # 添加中心节点
    dot.node('C', filename)

    #处理子公司性质和股东名字,加上换行符,使得显示不那么拥挤
    Subsidiary_characteristics_modified_data = Subsidiary_characteristics.apply(insert_newline_every_4_chars)

    shareHolder_names = data_list_after_combine[0].iloc[:,0].apply(insert_newline_every_2_chars)
    shareHolder_characteristics = data_list_after_combine[0].iloc[:,1]

    # 画股东表
    shareHolder_names[10] = '凑数'
    for i in range(len(shareHolder_names)):
        if i != len(shareHolder_names)-1:
            if '境外' in shareHolder_characteristics[i]:
                dot.node(f'S{i}', shareHolder_names.loc[i], color = 'red', height='2',width='1')
            else:
                dot.node(f'S{i}', shareHolder_names.loc[i], height='2',width='1')
            dot.node(f'HS{i}','HiddenPoint',shape='point',width='0')
            dot.edge(f'S{i}',f'HS{i}',label=data_list_after_combine[0].iloc[i,2],arrowhead='none')
            if i != 0:
                dot.edge(f'HS{i-1}',f'HS{i}',constraint='false',arrowhead='none')
        else:
            dot.node(f'S{i}', shareHolder_names[i],style='invis', height='2',width='1')
            dot.node(f'HS{i}','HiddenPoint',shape='point',width='0',style='invis')
            dot.edge(f'S{i}',f'HS{i}',style='invis')
            if i != 0:
                dot.edge(f'HS{i-1}',f'HS{i}',constraint='false',style='invis',arrowhead='none')
    #让公司名称的指向保留最中间的一个edge
    for i in range(len(shareHolder_names)):
        if i != (len(shareHolder_names)-1)/2:
            dot.edge(f'HS{i}','C',arrowhead='none',style='invis')
        else:
            dot.edge(f'HS{i}','C',arrowhead='none')


    #画实际控制人
    actual_controller = ''
    # 合成实控人字符串
    for i in range(len(name_list)):
        actual_controller = f'{actual_controller}   {name_list[i]}'

    # 将实控人字符串显示在图中
    dot.node('AC', actual_controller)
    dot.edge('AC', 'S0', style='invis')


    # 画子公司性质
    # 判断子公司性质是否为偶数
    if len(Subsidiary_characteristics_modified_data)%2 != 1:
        Subsidiary_characteristics_modified_data[len(Subsidiary_characteristics_modified_data)] = '凑数dede'
        for i in range(len(Subsidiary_characteristics_modified_data)):
            if i != len(Subsidiary_characteristics_modified_data)-1:
                dot.node(f'Scm{i}', Subsidiary_characteristics_modified_data[i], height='1',width='3')
                dot.node(f'HC{i}','HiddenPoint',shape='point',width='0')
                dot.edge(f'HC{i}',f'Scm{i}')
                if i != 0:
                    dot.edge(f'HC{i-1}',f'HC{i}',constraint='false',arrowhead='none')
            else:
                dot.node(f'Scm{i}', Subsidiary_characteristics_modified_data[i],style='invis', height='1',width='3')
                dot.node(f'HC{i}','HiddenPoint',shape='point',width='0',style='invis')
                dot.edge(f'HC{i}',f'Scm{i}',style='invis')
                if i != 0:
                    dot.edge(f'HC{i-1}',f'HC{i}',constraint='false',style='invis',arrowhead='none')
    # 奇数的情况
    else:
        for i in range(len(Subsidiary_characteristics_modified_data)):
            dot.node(f'Scm{i}', Subsidiary_characteristics_modified_data[i], height='1',width='3')
            dot.node(f'HC{i}','HiddenPoint',shape='point',width='0')
            dot.edge(f'HC{i}',f'Scm{i}')
            if i != 0:
                dot.edge(f'HC{i-1}',f'HC{i}',constraint='false',arrowhead='none')
    # 让公司名称的指向保留最中间的一个edge
    for i in range(len(Subsidiary_characteristics_modified_data)):
        if i != (len(Subsidiary_characteristics_modified_data)-1)/2:
            dot.edge('C',f'HC{i}',arrowhead='none',style='invis')
        else:
            dot.edge('C',f'HC{i}',arrowhead='none')

    # 加载JSON格式的字典
    with open('models_Judge_regions_of_composed_comanys/extended_dict_of_judge_Regions_model.json', 'r') as json_file:
        vocab = json.load(json_file)

    # 首先，重新创建模型架构
    model = SimpleClassifier_for_Overseas_Entities(vocab_size=len(vocab), embedding_dim=4, hidden_dim=8, output_dim=2)  # 替换MyModel为你的模型类

    # 然后，加载模型状态
    model.load_state_dict(torch.load('models_Judge_regions_of_composed_comanys/judge_Regions_of_Operation_of_composedcomp.pth'))
    # #尝试归类第一个性质
    j = 0
    for characteristic in Subsidiary_characteristics:
        
        bool_series_for_same_characteristic = data_list_after_combine[1].iloc[:,3] == characteristic
        # 获取满足条件的索引
        true_indices = bool_series_for_same_characteristic.index[bool_series_for_same_characteristic]
        i = 0

        all_subsidiary_names = ''

        overseas = ''

        for index in true_indices:
            subsidiary_name = data_list_after_combine[1].iloc[index, 0]
            Ownership_Stake_in_Subsidiary = data_list_after_combine[1].iloc[index, 4]
            register_place = data_list_after_combine[1].iloc[index, 2]
            # 禁用梯度计算
            # 将新文本数据转换为数值表示
            new_sequence = text_to_indices(register_place, vocab)
            predicted = torch.Tensor()
            with torch.no_grad():
                # 转换为适合模型的输入格式
                inputs = torch.tensor([new_sequence])
                
                # 获取模型预测
                outputs = model(inputs)

                # 选择概率最高的类别作为预测结果
                _, predicted = torch.max(outputs, 1)

            if predicted.item() == 0:
                all_subsidiary_names = f'{all_subsidiary_names} {subsidiary_name} {str(Ownership_Stake_in_Subsidiary)}%\n'
            else: 
                overseas = f'{overseas} {subsidiary_name} {str(Ownership_Stake_in_Subsidiary)}%\n'
            i+=1

        if all_subsidiary_names:
            dot.node(f'RH{characteristic}_all_subsidiary_names', f'境内子公司\n\n{all_subsidiary_names}', width='2')
            dot.edge(f'Scm{j}', f'RH{characteristic}_all_subsidiary_names')
        else:
            dot.node(f'RH{characteristic}_all_subsidiary_names', '无境内子公司', width='2')
            dot.edge(f'Scm{j}', f'RH{characteristic}_all_subsidiary_names')

        if overseas:
            dot.node(f'RH{characteristic}_overseas', f'境外子公司\n\n{overseas}', width='2')
            dot.edge(f'RH{characteristic}_all_subsidiary_names', f'RH{characteristic}_overseas',style='invis')
        else:
            dot.node(f'RH{characteristic}_overseas', '无境外子公司', width='2')
            dot.edge(f'RH{characteristic}_all_subsidiary_names', f'RH{characteristic}_overseas',style='invis')
        j+=1

    # 保存并查看图像
    dot.render(f'svg_graphs/{filename[:-5]}', format='svg',cleanup=True)

# 定义神经网络模型
class SimpleClassifier_for_Overseas_Entities(nn.Module):
    def __init__(self, vocab_size, embedding_dim, hidden_dim, output_dim):
        super(SimpleClassifier_for_Overseas_Entities, self).__init__()
        self.embedding = nn.Embedding(num_embeddings=vocab_size, embedding_dim=embedding_dim)
        self.fc1 = nn.Linear(embedding_dim, hidden_dim)
        self.fc2 = nn.Linear(hidden_dim, output_dim)

    def forward(self, x):
        x = self.embedding(x).sum(dim=1)
        x = F.relu(self.fc1(x))
        x = self.fc2(x)
        return x
# 创建自定义数据集
class CityDataset(Dataset):
    def __init__(self, sequences, labels):
        self.sequences = sequences
        self.labels = labels

    def __len__(self):
        return len(self.sequences)

    def __getitem__(self, idx):
        return torch.tensor(self.sequences[idx]), torch.tensor(self.labels[idx])

def draw_figure(filepath):
    concat_tables(filepath)
    print(f"failed to concat all tables")
    for filename in os.listdir('xlsx包括合成好的三张表'):
        try:
            data_list_after_combine, Subsidiary_characteristics, name_list = get_three_tables_and_namelist(filename, f'xlsx包括合成好的三张表/{filename}')
            draw_svg(data_list_after_combine, Subsidiary_characteristics, name_list, filename)
        except Exception as e:  # 捕获所有异常
            print(f"(rare_form_of_the_report,or Model's judgment is wrong)An error occurred while processing {filename}: {e}")


def modify_svg(filepath):
    for filename in os.listdir(filepath):
        svg_manipulator = SVGManipulator(f'{filepath}/{filename}')
        js_code = svg_manipulator.update_svg_script()  # 生成JS代码
        if js_code:  # 如果成功生成了JS代码
            svg_manipulator.insert_js_into_svg(js_code)  # 将JS代码插入SVG文件

def Judge_and_draw(filepath='test_word'):
    # dir_name = "your_directory_name"
    # # 检查当前目录下是否存在该文件夹
    # if os.path.exists(dir_name) and os.path.isdir(dir_name):
    #     # 删除这个目录及其所有内容
    #     shutil.rmtree(dir_name)
    #     print(f"The directory '{dir_name}' and all its contents have been removed.")
    # else:
    #     print(f"The directory '{dir_name}' does not exist.")
    Judge_fun(filepath)
    draw_figure('xlsx被模型标记完成的')
    if not os.path.exists('svg_graphs'):
        os.mkdir('svg_graphs')
    modify_svg('svg_graphs')
    

#训练代码
#打开docx转换为xlsx
def extract_table_in_docx(filepath):
    for filename in os.listdir(filepath):
        doc = Document(f'{filepath}/{filename}')  # 替换为你的文件名
        dataframes = []  # 存储所有表格的DataFrame列表

        elements = []  # 存储文档的所有元素
        # 遍历文档中的所有元素，收集段落和表格
        table_counter = 0
        for element in doc.element.body:
            if element.tag.endswith('p'):
                elements.append(('paragraph', element.text.strip()))
            elif element.tag.endswith('tbl'):
                elements.append(('table', doc.tables[table_counter]))
                table_counter+=1
        # 处理每个表格
        for i, (etype, content) in enumerate(elements):
            if etype == 'table':
                # 初始化用于存储当前表格数据的列表
                table_data = []
                # 初始化用于存储将添加到表格顶部的段落文本的列表
                preceding_paragraphs = []
                # 向上回溯查找直到找到三个非空文本段落或遇到一个表格
                count = 0  # 用于计数找到的非空文本段落
                for prev_index in range(i - 1, -1, -1):  # 从当前位置向上遍历
                    prev_element = elements[prev_index]
                    if prev_element[0] == 'paragraph' and prev_element[1]:  # 是非空段落
                        preceding_paragraphs.insert(0, prev_element[1])  # 将段落文本插入列表的开始

                        count += 1
                        if count == 5:  # 如果已找到三个非空文本段落，则jkuio89止
                            break
                    elif prev_element[0] == 'table':  # 如果遇到另一个表格，则停止
                        if count == 0:
                            preceding_paragraphs.insert(0, '无上信息')
                        break

                # 将找到的段落作为新行添加到表格顶部
                for para_text in preceding_paragraphs:
                    paragraph_row = [para_text] + [None]*(len(content.rows[0].cells)-1)  # 填充段落文本到整行
                    table_data.append(paragraph_row)

                # 添加表格的其他行
                for row in content.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                    break

                # 将处理好的表格数据转换成DataFrame并存储
                df = pd.DataFrame(table_data)
                dataframes.append(df)
        # 创建文件夹
        if not os.path.exists('xlsx前五行信息'):
            os.makedirs('xlsx前五行信息')
        # 创建一个Excel writer对象，使用pandas的ExcelWriter
        with pd.ExcelWriter(f'xlsx前五行信息/{filename[0:-5]}.xlsx') as writer:
            for i, df in enumerate(dataframes):
                df.to_excel(writer,sheet_name=f'Sheet{i+1}')
def read_labeled_xlsx2df(filepath):
    dataframes_for_train = []
    labels_for_train = []
    for filename in os.listdir(filepath):
        
        # 打开Excel文件
        workbook = openpyxl.load_workbook(f'{filepath}/{filename}')
        # 获取所有工作表的名称
        sheet_names = workbook.sheetnames
        # 遍历所有工作表
        for sheet_name in sheet_names:
            if '股东' in sheet_name:
                df = pd.read_excel(f'{filepath}/{filename}', sheet_name=sheet_name,index_col=0,header=0,engine='openpyxl')
                dataframes_for_train.append(pd.DataFrame(df.values))
                labels_for_train.append(1)
            elif '子公司' in sheet_name:
                df = pd.read_excel(f'{filepath}/{filename}', sheet_name=sheet_name, index_col=0, header=0, engine='openpyxl')
                dataframes_for_train.append(pd.DataFrame(df.values))
                labels_for_train.append(2)
            elif '实际控制人' in sheet_name:
                df = pd.read_excel(f'{filepath}/{filename}', sheet_name=sheet_name, index_col=0, header=0, engine='openpyxl')
                dataframes_for_train.append(pd.DataFrame(df.values))
                labels_for_train.append(3)
            else:
                df = pd.read_excel(f'{filepath}/{filename}', sheet_name=sheet_name, index_col=0, header=0, engine='openpyxl')
                if df.iloc[0,0] == '无上信息':
                    continue
                dataframes_for_train.append(pd.DataFrame(df.values))
                labels_for_train.append(4)
        workbook.close()
        
    for i, df in enumerate(dataframes_for_train):
        dataframes_for_train[i] = pd.DataFrame(df)
    return dataframes_for_train, labels_for_train
# 将表格转换为文本
def table_to_text_for_one_report(table):
    train_text_without_none = [x for x in [' '.join(str(cell) for cell in row[::-1] if pd.isnull(cell) != True) for row in table.values[::-1]]]
    return ' '.join(train_text_without_none)
#只拿出包含股东,实控人, 子公司的表格作为训练集
def select_tables_and_conv2texts(df_list_for_judge, new_labels_af_select):

    # 准备数据
    all_texts = [table_to_text_for_one_report(data) for data in df_list_for_judge]
    # 要查找的字符
    chars_to_check = ["实际控制人", "股东",'子公司']

    # 使用列表推导式检查每个元素是否包含任何一个指定字符，并将包含字符的元素添加到新列表中
    all_texts_af_del = [item for item in all_texts if any(char in item for char in chars_to_check)]
    # 使用列表推导式获取包含指定字符的元素的索引
    indices = [label for label, item in zip(new_labels_af_select,all_texts) if any(char in item for char in chars_to_check)]
    # 确定类别4的索引
    indexes_of_4 = [index for index, value in enumerate(indices) if value == 4]
    indexes_of_4_af_del = indexes_of_4
    # 计算非4的总数和所需的4的数量
    non_4_count = len([label for label in indices if label != 4])
    required_4_count = non_4_count  # 因为我们希望4的数量等于非4的总数
    texts_af_average = []
    labels_af_average = []
    # 如果类别4的样本多于所需数量，则减少到所需数量
    if len(indexes_of_4) > required_4_count:
        # 随机选择所需数量的类别4的索引
        # 把邻近的两个表格加入到需要训练的类别四中
        
        retained_indexes_of_4 = random.sample(indexes_of_4_af_del, required_4_count)
        # 添加所需数量的类别4的样本到新列表
        for index in retained_indexes_of_4:
            texts_af_average.append(all_texts_af_del[index])
            labels_af_average.append(indices[index])

    # 添加所有非4的样本到新列表
    for index, label in enumerate(indices):
        if label != 4:
            texts_af_average.append(all_texts_af_del[index])
            labels_af_average.append(label)


    return texts_af_average, labels_af_average

def remove_punctuation(text):
    # 使用正则表达式替换掉文本中的标点符号
    punctuation = """ )(！？｡。＂＃＄％＆＇（）＊＋，－／：；＜＝＞＠［＼］＾＿｀｛｜｝～｟｠｢｣､、〃《》「」『』【】〔〕〖〗〘〙〚〛〜〝□〞〟〰〾〿√–—‘’‛“”„‟…‧﹏.\n\r\t"""
    return re.sub(f"[{punctuation}]", "", text)

def split_data_and_save_to_csv(text_sequences_list, labels_af_del):
    # 处理列表中的每个字符串
    clean_texts = [remove_punctuation(text) for text in text_sequences_list]

    # 假设 X 是你的特征（例如文本数据），y 是对应的标签
    X_train, X_valid, y_train, y_valid = train_test_split(clean_texts, labels_af_del, test_size=0.2, random_state=42)

    # 打开一个新的CSV文件，用来写入训练集
    with open('texts_labels_withkeywords_reversed.csv', mode='w', newline='', encoding='utf-8',) as file:
        writer = csv.writer(file,delimiter='|')
        # 写入数据
        for text, label in zip(X_train, y_train):
            writer.writerow([text, label-1])
    # 打开一个新的CSV文件, 用来写入验证集
    with open('texts_labels_withkeywords_tests_reversed.csv', mode='w', newline='', encoding='utf-8',) as file:
        writer = csv.writer(file, delimiter='|')
        # 写入数据
        for text, label in zip(X_valid, y_valid):
            writer.writerow([text, label-1])

def create_train_and_test_set(filepath='xlsx前五行信息'):
    dataframes_for_train, labels_for_train = read_labeled_xlsx2df(filepath)
    text_sequences_list, labels_af_del = select_tables_and_conv2texts(dataframes_for_train, labels_for_train)
    split_data_and_save_to_csv(text_sequences_list, labels_af_del)


def train(model, dataloader, optimizer, tokenizer, device):
    model.train()  # 设置模型为训练模式
    epoch_loss = 0
    epoch_acc = 0
    for i, batch in enumerate(dataloader):
        label = batch["label"].to(device)  # 将标签移动到指定设备
        text = batch["text"]  # 获取文本数据

        # 对文本进行编码处理
        tokenized_text = tokenizer(text, max_length=100, add_special_tokens=True, truncation=True, padding=True, return_tensors="pt")
        tokenized_text = {k: v.to(device) for k, v in tokenized_text.items()}  # 移动到指定设备

        optimizer.zero_grad()  # 清除旧的梯度

        # 进行一次前向传播
        output = model(**tokenized_text, labels=label)
        loss = output.loss  # 损失
        logits = output.logits  # 模型预测的原始输出

        # 计算准确率
        y_pred_label = torch.argmax(logits, dim=1)  # 从logits获取预测的标签
        acc = (y_pred_label == label).float().mean()  # 计算准确率

        # 反向传播和优化
        loss.backward()
        optimizer.step()

        # 累加本epoch的损失和准确率
        epoch_loss += loss.item()
        epoch_acc += acc.item()

        # 每200批次打印一次当前损失和准确率
        if i % 200 == 0:
            print(f"current loss: {epoch_loss / (i+1):.4f} \t current acc: {epoch_acc / (i+1):.4f}")

    # 返回本epoch的平均损失和准确率
    return epoch_loss / len(dataloader), epoch_acc / len(dataloader)

def evaluate(model, dataloader, tokenizer, device):
    model.eval()
    epoch_loss = 0
    epoch_acc = 0
    with torch.no_grad():
        for _, batch in enumerate(dataloader):
            label = batch["label"].to(device) 
            text = batch["text"]
            tokenized_text = tokenizer(text, max_length=100, add_special_tokens=True, truncation=True, padding=True,
                                    return_tensors="pt")
            tokenized_text = tokenized_text.to(device)

            # 进行一次前向传播
            output = model(**tokenized_text, labels=label)
            loss = output.loss  # 损失
            logits = output.logits  # 模型预测的原始输出

            # 计算准确率
            y_pred_label = torch.argmax(logits, dim=1)  # 从logits获取预测的标签
            acc = (y_pred_label == label).float().mean()  # 计算准确率

            # 累加本epoch的损失和准确率
            epoch_loss += loss.item()
            epoch_acc += acc.item()

    # len(dataloader) 表示有多少个 batch，len(dataloader.dataset.dataset) 表示样本数量
    return epoch_loss / len(dataloader), epoch_acc / len(dataloader.dataset.dataset)

class SentimentDataset(Dataset):
    def __init__(self, path_to_file):
        self.dataset = pd.read_csv(path_to_file, names=["text", "label"], sep='|')
    def __len__(self):
        return len(self.dataset)
    def __getitem__(self, idx):
        # 根据 idx 分别找到 text 和 label
        text = self.dataset.loc[idx, "text"]
        label = self.dataset.loc[idx, "label"]
        # 将label从字符串转换为整数
        label = int(label)
        label_tensor = torch.tensor(label, dtype=torch.long)
        sample = {"text": text, "label": label_tensor }
        # 返回一个 dict
        return sample

def train_for_finding_tables(train_set_path='texts_labels_withkeywords_reversed.csv',test_set_path='texts_labels_withkeywords_tests_reversed.csv',epochs=10):
    global global_model_version
    global_model_version+=1
    os.environ['CUDA_VISIBLE_DEVICES'] = '0'

    # 超参数
    hidden_dropout_prob = 0.3
    num_labels = 4
    learning_rate = 2e-5
    weight_decay = 1e-2

    batch_size = 1
    device = 'cpu' if not torch.cuda.is_available() else 'cuda'

    vocab_file = "models_find_tables/vocab.txt" # 词汇表


            
    # 加载训练集
    sentiment_train_set = SentimentDataset(train_set_path)
    sentiment_train_loader = DataLoader(sentiment_train_set, batch_size=batch_size, shuffle=True, num_workers=0)
    # 加载验证集
    sentiment_valid_set = SentimentDataset(test_set_path)
    sentiment_valid_loader = DataLoader(sentiment_valid_set, batch_size=batch_size, shuffle=False, num_workers=0)



    # 定义 tokenizer，传入词汇表
    tokenizer = BertTokenizer(vocab_file)


    # 加载模型
    config = BertConfig.from_pretrained("bert-base-uncased", num_labels=num_labels, hidden_dropout_prob=hidden_dropout_prob)
    model = BertForSequenceClassification.from_pretrained("bert-base-uncased", config=config)
    model.to(device)


    # 定义优化器和损失函数
    # Prepare optimizer and schedule (linear warmup and decay)
    # 设置 bias 和 LayerNorm.weight 不使用 weight_decay
    no_decay = ['bias', 'LayerNorm.weight']
    optimizer_grouped_parameters = [
            {'params': [p for n, p in model.named_parameters() if not any(nd in n for nd in no_decay)], 'weight_decay': weight_decay},
            {'params': [p for n, p in model.named_parameters() if any(nd in n for nd in no_decay)], 'weight_decay': 0.0}
    ]

    #optimizer = AdamW(model.parameters(), lr=learning_rate)
    optimizer = AdamW(optimizer_grouped_parameters, lr=learning_rate)
    criterion = nn.CrossEntropyLoss()
        # 开始训练和验证
    for i in range(epochs):
        train_loss, train_acc = train(model, sentiment_train_loader, optimizer, tokenizer, device)
        print("train loss: ", train_loss, "\t", "train acc:", train_acc)
        valid_loss, valid_acc = evaluate(model, sentiment_valid_loader,tokenizer, device)
        print("valid loss: ", valid_loss, "\t", "valid acc:", valid_acc)

    # 保存模型
    if not os.path.exists('models_find_tables'):
        os.mkdir('models_find_tables')
    torch.save(model.state_dict(), f'models_find_tables/trained_model_bertclassifier_epoch_5_keywords_reversed{global_model_version}.pth')


def text_to_sequence(text, char_to_index):
    return [char_to_index[char] for char in text]

# 填充文本序列
def pad_sequence(seq, max_length):
    seq += [0] * (max_length - len(seq))  # 使用0作为填充值
    return seq

def train_for_identify_Overseas_Entities(num_epochs=70):
    tokenizer = BertTokenizer.from_pretrained('bert-base-chinese')  # 使用中文BERT模型的名称
    vocab = tokenizer.get_vocab()  # 获取词汇表字典
    chinese_city = '''福耀, 太仓, 哈尔滨, 齐齐哈尔, 牡丹江, 佳木斯, 大庆, 伊春, 鸡西, 鹤岗, 双鸭山, 七台河, 黑河, 绥化, 大兴安岭地区, 拉萨, 日喀则, 林芝, 昌都, 那曲, 阿里, 山南, 乌鲁木齐, 克拉玛依, 喀什, 伊犁, 阿克苏, 塔城, 哈密, 昌吉, 阿勒泰, 吐鲁番, 巴音郭楞, 克孜勒苏, 和田, 霍尔果斯, 北京, 上海, 天津, 重庆, 石家庄, 郑州, 武汉, 长沙, 南京, 杭州, 合肥, 福州, 南昌, 济南, 太原, 西安, 兰州, 西宁, 成都, 贵阳, 昆明, 拉萨, 广州, 海口, 南宁, 哈尔滨, 长春, 沈阳, 呼和浩特, 乌鲁木齐, 银川, 青岛, 大连, 宁波, 厦门, 深圳, 苏州, 无锡, 杭州, 成都, 重庆, 武汉, 西安, 长沙, 南京, 广州, 深圳, 上海, 北京, 唐山, 秦皇岛, 邯郸, 邢台, 保定, 张家口, 承德, 沧州, 廊坊, 衡水, 洛阳, 平顶山, 安阳, 鹤壁, 新乡, 焦作, 濮阳, 许昌, 漯河, 三门峡, 南阳, 商丘, 信阳, 周口, 驻马店, 黄石, 十堰, 宜昌, 襄阳, 鄂州, 荆门, 孝感, 荆州, 黄冈, 咸宁, 随州, 恩施土家族苗族自治州, 株洲, 湘潭, 衡阳, 邵阳, 岳阳, 常德, 张家界, 益阳, 郴州, 永州, 怀化, 娄底, 湘西土家族苗族自治州, 徐州, 常州, 泰州, 宿迁, 盐城, 扬州, 镇江, 泰州, 宿迁, 湛江, 茂名, 肇庆, 珠海, 佛山, 江门, 韶关, 惠州, 梅州, 汕头, 深圳, 珠海, 佛山, 江门, 潮州, 揭阳, 中山, 东莞, 清远, 阳江, 云浮, 南宁, 柳州, 桂林, 梧州, 北海, 防城港, 钦州, 贵港, 玉林, 百色, 贺州, 河池, 来宾, 崇左, 海口, 三亚, 三沙, 儋州, 五指山, 文昌, 万宁, 东方, 定安县, 屯昌县, 澄迈县, 临高县, 白沙黎族自治县, 昌江黎族自治县, 乐东黎族自治县, 陵水黎族自治县, 保亭黎族苗族自治县, 琼中黎族苗族自治县, 西双版纳傣族自治州, 德宏傣族景颇族自治州, 大理白族自治州, 怒江傈僳族自治州, 迪庆藏族自治州, 楚雄彝族自治州, 红河哈尼族彝族自治州, 文山壮族苗族自治州, 普洱, 临沧, 保山, 昭通, 曲靖, 玉溪, 丽江, 昆明, 普洱, 临沧, 保山, 昭通, 曲靖, 玉溪, 丽江, 绵阳, 自贡, 攀枝花, 泸州, 德阳, 广元, 遂宁, 内江, 乐山, 南充, 眉山, 宜宾, 广安, 达州, 雅安, 巴中, 资阳, 阿坝藏族羌族自治州, 甘孜藏族自治州, 凉山彝族自治州, 贵阳, 六盘水, 遵义, 安顺, 毕节, 铜仁, 黔西南布依族苗族自治州, 黔东南苗族侗族自治州, 黔南布依族苗族自治州, 曲靖, 玉溪, 保山, 昭通, 丽江, 普洱, 临沧, 楚雄彝族自治州, 红河哈尼族彝族自治州, 文山壮族苗族自治州, 西双版纳傣族自治州, 大理白族自治州, 德宏傣族景颇族自治州, 怒江傈僳族自治州, 迪庆藏族自治州, 兰州, 嘉峪关, 金昌, 白银, 天水, 武威, 张掖, 平凉, 酒泉, 庆阳, 定西, 陇南, 临夏回族自治州, 甘南藏族自治州, 西宁, 海东, 海北藏族自治州, 黄南藏族自治州, 海南藏族自治州, 果洛藏族自治州, 玉树藏族自治州, 海西蒙古族藏族自治州, 银川, 石嘴山, 吴忠, 固原, 中卫, 乌鲁木齐, 克拉玛依, 吐鲁番, 哈密, 昌吉回族自治州, 博尔塔拉蒙古自治州, 巴音郭楞蒙古自治州, 阿克苏地区, 克孜勒苏柯尔克孜自治州, 喀什地区, 和田地区, 伊犁哈萨克自治州, 塔城地区, 阿勒泰地区, 石河子, 阿拉尔, 图木舒克, 五家渠, 北屯, 双河, 可克达拉, 昆玉, 胡杨河'''
    out_of_chinese = '''美国， 日本， 德国， 英国， 法国， 俄罗斯， 印度， 巴西， 加拿大， 澳大利亚， 意大利， 西班牙， 墨西哥， 印度尼西亚， 土耳其， 荷兰， 波兰， 泰国， 瑞士， 瑞典， 挪威， 奥地利， 芬兰， 丹麦， 以色列， 葡萄牙， 爱尔兰， 南非， 阿根廷， 哥伦比亚， 智利， 马来西亚， 菲律宾， 新加坡， 越南， 新西兰， 比利时， 沙特阿拉伯， 伊朗， 韩国， 巴基斯坦， 孟加拉国， 埃及， 阿尔及利亚， 尼日利亚， 肯尼亚， 埃塞俄比亚， 摩洛哥， 突尼斯， 斯里兰卡， 阿富汗， 安哥拉， 贝宁， 布基纳法索， 布隆迪， 喀麦隆， 佛得角， 中非共和国， 乍得， 科特迪瓦， 刚果民主共和国， 刚果共和国， 吉布提， 厄立特里亚， 加蓬， 冈比亚， 加纳， 几内亚， 几内亚比绍， 象牙海岸， 利比里亚， 马达加斯加， 马拉维， 马里， 毛里塔尼亚， 莫桑比克， 纳米比亚， 尼日尔， 卢旺达， 塞内加尔， 塞拉利昂， 索马里， 南苏丹， 苏丹， 坦桑尼亚， 多哥， 乌干达， 赞比亚， 津巴布韦'''
    chinese_city_list = chinese_city.replace(' ', '').split(',')
    out_of_chinese_list = out_of_chinese.replace(' ', '').split('，')
    cities = chinese_city_list + out_of_chinese_list
    #定义列表长度
    length1 = len(out_of_chinese_list)
    length2 = len(chinese_city_list)
    zero_list = [0] * length2
    one_list = [1] * length1
    labels = zero_list + one_list
    char_list = list(set(','.join(cities)))
    # 检查自定义词汇是否已存在于BERT的词汇表中，并删除重合的部分
    existing_words = set(char_list).intersection(vocab.keys())
    # 删除重合的部分
    char_list_after_delete = list(set(char_list) - existing_words)
    for word in char_list_after_delete:
        vocab[word] = len(vocab)  # 分配新的索引
    sequences = [text_to_sequence(city, vocab) for city in cities]
    # 找到最大长度的序列
    max_length = max(len(seq) for seq in sequences)  
    # 应用填充
    padded_sequences = [pad_sequence(seq, max_length) for seq in sequences]
    dataset = CityDataset(padded_sequences, labels)
    dataloader = DataLoader(dataset, batch_size=1, shuffle=True)
    # 实例化模型
    model = SimpleClassifier_for_Overseas_Entities(vocab_size=len(vocab), embedding_dim=4, hidden_dim=8, output_dim=2)
    # 损失函数和优化器
    criterion = nn.CrossEntropyLoss()
    optimizer = optim.Adam(model.parameters(), lr=0.001)
    # 训练模型
    for epoch in range(num_epochs):
        for inputs, targets in dataloader:
            # 前向传播
            outputs = model(inputs)
            loss = criterion(outputs, targets)
            # 反向传播和优化
            optimizer.zero_grad()
            loss.backward()
            optimizer.step()
        print(f'Epoch {epoch+1}/{num_epochs}, Loss: {loss.item():.4f}')
    if not os.path.exists('models_Judge_regions_of_composed_comanys'):
        os.mkdir('models_Judge_regions_of_composed_comanys')
    # 选择一个文件名
    dict_file = 'models_Judge_regions_of_composed_comanys/extended_dict_of_judge_Regions_model.json'
    # 将字典写入文件
    with open(dict_file, 'w') as file:
        json.dump(vocab, file)
    # 选择一个文件名
    model_file = 'models_Judge_regions_of_composed_comanys/judge_Regions_of_Operation_of_composedcomp.pth'
    # 保存模型
    torch.save(model.state_dict(), model_file)


class SVGManipulator:
    def __init__(self, svg_file_path):
        self.svg_file_path = svg_file_path

    def find_max_scm_value(self):
        with open(self.svg_file_path, 'r', encoding='utf-8') as file:
            svg_content = file.read()
        scm_comments = re.findall(r'<!-- Scm(\d+) -->', svg_content)
        max_scm = max(map(int, scm_comments), default=-1)
        return max_scm

    def count_foreign_subsidiaries(self):
        with open(self.svg_file_path, 'r', encoding='utf-8') as file:
            svg_content = file.read()
        count = len(re.findall(r'境外子公司', svg_content))
        return count

    def generate_subcategories_even(self, max_scm):
        sub_categories = {}
        last_scm_id = 25 + (max_scm - 1) * 2  # 因为我们从0开始计数
        first_subsidiary_id = last_scm_id + 4  # 子公司ID从最后一个SCM节点ID加4开始
        for i in range(max_scm):
            scm_id = 25 + i * 2
            subsidiaries_ids = [f'node{first_subsidiary_id + 2 * i}', f'node{first_subsidiary_id + 2 * i + 1}']
            sub_categories[f'node{scm_id}'] = subsidiaries_ids
        return sub_categories

    def generate_subcategories_odd(self, max_scm):
        sub_categories = {}
        last_scm_id = 25 + (max_scm) * 2  # 因为我们从0开始计数
        first_subsidiary_id = last_scm_id + 2  # 子公司ID从最后一个SCM节点ID加2开始
        for i in range(max_scm + 1):
            scm_id = 25 + i * 2
            subsidiaries_ids = [f'node{first_subsidiary_id + 2 * i}', f'node{first_subsidiary_id + 2 * i + 1}']
            sub_categories[f'node{scm_id}'] = subsidiaries_ids
        return sub_categories

    def update_svg_script(self):
        max_scm = self.find_max_scm_value()
        foreign_subs_count = self.count_foreign_subsidiaries()
        if foreign_subs_count % 2 == 0:  # 偶数情况
            sub_categories = self.generate_subcategories_even(max_scm)
        else:  # 奇数情况
            sub_categories = self.generate_subcategories_odd(max_scm)
        
        static_sub_nodes_ids = ['node2', 'node4', 'node6', 'node8', 'node10', 'node12', 'node14', 'node16', 'node18', 'node20']
        scm_nodes_ids = [f'node{25 + i * 2}' for i in range(max_scm + (1 if foreign_subs_count % 2 else 0))]
        main_sub_nodes = static_sub_nodes_ids + scm_nodes_ids

        # 构建 JavaScript 代码
        sub_categories_str = ', '.join([f"'{k}': ['{v[0]}', '{v[1]}']" for k, v in sub_categories.items()])
        main_sub_nodes_str = ', '.join([f"'{n}'" for n in main_sub_nodes])

        js_script = f"""
        <script type="text/javascript"><![CDATA[
            window.onload = function() {{
                function toggleNodes(nodeIds) {{
                    nodeIds.forEach(function(id) {{
                        var element = document.getElementById(id);
                        if (element) {{
                            element.style.display = element.style.display === 'none' ? '' : 'none';
                        }}
                    }});
                }}
                var allNodesExceptMain = document.querySelectorAll('g.node:not(#node1)');
                allNodesExceptMain.forEach(function(node) {{
                    node.style.display = 'none';
                }});
                var mainCategory = 'node1';
                var mainSubNodes = [{main_sub_nodes_str}];
                var subCategories = {{{sub_categories_str}}};
                var mainNode = document.getElementById(mainCategory);
                if (mainNode) {{
                    mainNode.addEventListener('click', function() {{
                        toggleNodes(mainSubNodes);
                    }});
                }}
                Object.keys(subCategories).forEach(function(subCatId) {{
                    var subNode = document.getElementById(subCatId);
                    if (subNode) {{
                        subNode.addEventListener('click', function() {{
                            toggleNodes(subCategories[subCatId]);
                        }});
                    }}
                }});
                var specialNode = document.getElementById('node2');
                if (specialNode) {{
                    specialNode.addEventListener('click', function() {{
                        toggleNodes(['node24']);
                    }});
                }}
            }};
        ]]></script>
        """
        return js_script

    def insert_js_into_svg(self, js_code):
        with open(self.svg_file_path, 'r', encoding='utf-8') as file:
            svg_content = file.read()

        updated_svg_content = re.sub(r'(\</svg\>)', js_code + r'\1', svg_content, flags=re.IGNORECASE)

        with open(self.svg_file_path, 'w', encoding='utf-8') as file:
            file.write(updated_svg_content)